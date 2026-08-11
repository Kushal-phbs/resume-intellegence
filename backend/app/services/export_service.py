"""Document export service for tailored resume artifacts."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

from app.core.exceptions import ResourceNotFoundException, ValidationException
from app.models.cover_letter import CoverLetter
from app.models.resume_tailoring_version import ResumeTailoringVersion
from app.repositories.cover_letter_repository import CoverLetterRepository
from app.repositories.resume_repository import ResumeRepository
from app.repositories.resume_version_repository import ResumeVersionRepository
from app.storage.base import StorageProvider


class ExportService:
    """Builds export files for tailored resume versions and cover letters."""

    def __init__(
        self,
        resume_version_repository: ResumeVersionRepository,
        cover_letter_repository: CoverLetterRepository,
        resume_repository: ResumeRepository,
        storage_provider: StorageProvider,
    ) -> None:
        self._resume_versions = resume_version_repository
        self._cover_letters = cover_letter_repository
        self._resumes = resume_repository
        self._storage = storage_provider

    async def export_resume(
        self,
        *,
        user_id: UUID,
        version_id: UUID,
        format: str,
    ) -> Path:
        version = await self._resume_versions.get_by_id(version_id)
        if version is None:
            raise ResourceNotFoundException("Tailored resume version not found")

        await self._ensure_resume_owner(user_id=user_id, resume_id=version.resume_id)
        return self._write_export(
            content=self._render_resume_markdown(version),
            basename=f"tailored-resume-{version.id}",
            format=format,
        )

    async def export_cover_letter(
        self,
        *,
        user_id: UUID,
        cover_letter_id: UUID,
        format: str,
    ) -> Path:
        letter = await self._cover_letters.get_by_id(cover_letter_id)
        if letter is None:
            raise ResourceNotFoundException("Cover letter not found")

        await self._ensure_resume_owner(
            user_id=user_id,
            resume_id=letter.tailoring_session.resume_id,
        )
        return self._write_export(
            content=self._render_cover_letter_markdown(letter),
            basename=f"cover-letter-{letter.id}",
            format=format,
        )

    async def _ensure_resume_owner(self, *, user_id: UUID, resume_id: UUID) -> None:
        resume = await self._resumes.get(resume_id)
        if resume is None or resume.user_id != user_id:
            raise ResourceNotFoundException("Export not found")

    def _write_export(self, *, content: str, basename: str, format: str) -> Path:
        normalized = format.lower().strip()
        if normalized not in {"md", "docx", "pdf"}:
            raise ValidationException("Unsupported export format")

        if normalized == "md":
            extension = "md"
            raw_bytes = content.encode("utf-8")
        elif normalized == "docx":
            extension = "docx"
            raw_bytes = self._render_docx(content)
        elif normalized == "pdf":
            extension = "pdf"
            raw_bytes = self._render_pdf(content)
        else:
            raise ValidationException("Unsupported export format")

        storage_key = self._storage.save(
            content=raw_bytes,
            filename=f"{basename}-{uuid4()}.{extension}",
        )
        return self._storage.get_download_path(storage_key)

    def _render_docx(self, markdown_text: str) -> bytes:
        """Convert Markdown text to a valid OOXML DOCX using python-docx."""
        from docx import Document  # type: ignore[import-untyped]

        doc = Document()
        for line in markdown_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _render_pdf(self, markdown_text: str) -> bytes:
        """Convert Markdown text to a valid PDF using ReportLab."""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        styles = getSampleStyleSheet()
        story = []
        for line in markdown_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
            elif stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], styles["Title"]))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], styles["Heading2"]))
            elif stripped.startswith("- "):
                story.append(Paragraph(f"• {stripped[2:]}", styles["Normal"]))
            else:
                story.append(Paragraph(stripped, styles["Normal"]))
        doc.build(story)
        return buf.getvalue()

    def _render_resume_markdown(self, version: ResumeTailoringVersion) -> str:
        lines = [
            "# Tailored Resume",
            "",
            "## Professional Summary",
            version.professional_summary,
            "",
            "## Experience",
        ]
        if version.experience_json:
            for item in version.experience_json:
                lines.append(f"- {item}")
        else:
            lines.append("- No experience entries generated")

        lines.extend(["", "## Skills"])
        if version.skills_json:
            for item in version.skills_json:
                lines.append(f"- {item}")
        else:
            lines.append("- No skills entries generated")

        lines.extend(
            ["", "## ATS Score", str(version.ats_score), "", "## Recommendations"]
        )
        if version.recommendations_json:
            for item in version.recommendations_json:
                lines.append(f"- {item}")
        else:
            lines.append("- No recommendations generated")

        return "\n".join(lines)

    def _render_cover_letter_markdown(self, letter: CoverLetter) -> str:
        return "\n".join(
            [
                f"# {letter.title}",
                "",
                letter.greeting,
                "",
                letter.introduction,
                "",
                letter.body,
                "",
                letter.closing,
            ]
        )
